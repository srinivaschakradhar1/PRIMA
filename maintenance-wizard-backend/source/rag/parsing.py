"""Document parsing (layout- and structure-aware).

Extracts a canonical full-text string plus a list of :class:`DocumentSection`
objects whose ``start_offset`` / ``end_offset`` index into that string. Offsets
are preserved as the source of truth for all downstream chunking.

PDF parsing uses PyMuPDF (``fitz``) in *dict* mode so that font size, weight and
bounding boxes are available. That structural signal lets the parser:

* detect headings by typography (size / bold / all-caps) instead of fragile
  string heuristics, which previously promoted body sentences and clause numbers
  to headings;
* strip repeated running headers / footers (page numbers, document codes) that
  otherwise fragment a single logical section across page boundaries;
* reassemble headings that the PDF wrapped across several lines (e.g.
  ``"SPECIFICATION FOR CHAIN"`` + ``"WHEELS"`` -> one heading); and
* keep tables intact as a single block (rendered as pipe-delimited rows) instead
  of shredding them into one-line fragments.

Non-PDF inputs (txt, md, csv, docx, ...) and PDFs parsed without PyMuPDF fall
back to a line-based path that applies the same textual heading rules (markdown,
numbered, all-caps) without typographic signal. Word ``.docx`` files are first
flattened to text with python-docx, mapping their heading styles to markdown
``#`` prefixes so the same textual rules recover the document structure.
"""

from __future__ import annotations

import logging
import math
import re
import uuid
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from rag.config import SETTINGS
from rag.models import DocumentSection
from rag.tokenization import count_tokens

logger = logging.getLogger(__name__)

_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
_NUMBERED_HEADING = re.compile(r"^(\d+(?:\.\d+){0,5})[.)]?\s+(\S.{0,118})$")
_SENTENCE_END = (".", ":", ";", ",", "?", "!")
# Lines that look like a heading textually but are really figure/page markers.
_FIGURE_MARKER = re.compile(r"(?i)^(fig|figure|plate|photo)\.?\s*[\d.\-]*\s*$")
_PAGE_NUMBER = re.compile(r"(?i)^(page\s+)?\d{1,4}(\s*(of|/)\s*\d{1,4})?$")
_BOLD_FLAG = 1 << 4  # PyMuPDF span flag bit for bold/synthetic-bold faces.


@dataclass
class _Block:
    """A single text line or a rendered table, with layout metadata."""

    text: str
    page: int
    page_height: float
    y0: float
    size: float  # dominant font size (0.0 when unknown, e.g. plain text)
    bold: bool
    is_table: bool = False
    is_heading: bool = False
    level: int = 1


# ---------------------------------------------------------------------------
# Raw extraction
# ---------------------------------------------------------------------------
def _extract_pdf_blocks(file_path: Path) -> tuple[list[_Block], int] | None:
    """Return ``(blocks, num_pages)`` for a PDF, or ``None`` to use the fallback.

    Blocks are emitted in reading order; tables are interleaved at their vertical
    position. Headers/footers are *not* stripped here (that needs the whole-doc
    view); heading classification happens later.
    """
    try:  # pragma: no cover - optional dependency
        import fitz  # PyMuPDF
    except Exception:
        logger.warning("PyMuPDF (fitz) not installed; reading %s as raw text.", file_path.name)
        return None

    try:  # pragma: no cover - requires the dependency
        blocks: list[_Block] = []
        with fitz.open(file_path) as doc:
            num_pages = doc.page_count
            for page_index, page in enumerate(doc):
                page_no = page_index + 1
                page_height = float(page.rect.height) or 1.0
                table_rects = _table_rects(page)
                page_blocks: list[_Block] = []

                data = page.get_text("dict")
                for block in data.get("blocks", []):
                    if block.get("type") != 0:  # skip image blocks
                        continue
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        text = "".join(s.get("text", "") for s in spans)
                        if not text.strip():
                            continue
                        x0, y0, x1, y1 = line["bbox"]
                        cx, cy = (x0 + x1) / 2.0, (y0 + y1) / 2.0
                        if any(r.contains((cx, cy)) for r in table_rects):
                            continue  # belongs to a table; rendered separately
                        size = max((s.get("size", 0.0) for s in spans), default=0.0)
                        bold = any(
                            (s.get("flags", 0) & _BOLD_FLAG) or "bold" in s.get("font", "").lower()
                            for s in spans
                        )
                        page_blocks.append(
                            _Block(
                                text=text.strip(), page=page_no, page_height=page_height,
                                y0=float(y0), size=round(size, 1), bold=bold,
                            )
                        )

                for rect, rendered in _render_tables(page, table_rects):
                    if rendered.strip():
                        page_blocks.append(
                            _Block(
                                text=rendered, page=page_no, page_height=page_height,
                                y0=float(rect.y0), size=0.0, bold=False, is_table=True,
                            )
                        )

                page_blocks.sort(key=lambda b: b.y0)
                blocks.extend(page_blocks)
        return blocks, num_pages
    except Exception as exc:  # pragma: no cover
        logger.warning("PyMuPDF failed on %s (%s); reading as raw text.", file_path.name, exc)
        return None


def _table_rects(page):  # pragma: no cover - requires the dependency
    rects = []
    try:
        for tbl in page.find_tables().tables:
            if (tbl.row_count or 0) >= 2 and (tbl.col_count or 0) >= 2:
                rects.append(_Rect(tbl.bbox))
    except Exception:
        pass
    return rects


def _render_tables(page, table_rects):  # pragma: no cover - requires the dependency
    rendered = []
    try:
        for tbl in page.find_tables().tables:
            if (tbl.row_count or 0) < 2 or (tbl.col_count or 0) < 2:
                continue
            rendered.append((_Rect(tbl.bbox), _table_to_text(tbl)))
    except Exception:
        pass
    return rendered


def _table_to_text(tbl) -> str:  # pragma: no cover - requires the dependency
    try:
        rows = tbl.extract()
    except Exception:
        return ""
    lines = []
    for row in rows:
        cells = [" ".join((c or "").split()) for c in row]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


class _Rect:
    """Minimal axis-aligned rectangle (avoids importing fitz at module load)."""

    __slots__ = ("x0", "y0", "x1", "y1")

    def __init__(self, bbox) -> None:
        self.x0, self.y0, self.x1, self.y1 = (float(v) for v in bbox)

    def contains(self, point: tuple[float, float]) -> bool:
        x, y = point
        return self.x0 <= x <= self.x1 and self.y0 <= y <= self.y1


def _blocks_from_text(raw: str) -> list[_Block]:
    """Build single-line blocks from plain text (no typographic signal)."""
    blocks: list[_Block] = []
    for line in raw.splitlines():
        if line.strip():
            blocks.append(_Block(text=line.strip(), page=1, page_height=1.0, y0=0.0, size=0.0, bold=False))
    return blocks


_DOCX_HEADING_STYLE = re.compile(r"(?i)^heading\s+(\d+)$")


def _extract_text(file_path: Path) -> str:
    """Flatten a non-PDF document to plain text for the line-based parser.

    ``.docx`` files are unpacked with python-docx (heading styles become
    markdown ``#`` prefixes, tables become pipe-delimited rows). Plain-text
    formats (``.md``, ``.txt``, ``.csv``, ...) are already text and are decoded
    as UTF-8 with undecodable bytes dropped. A ``.docx`` that cannot be parsed
    falls back to the raw-bytes path so callers always get a string.
    """
    if file_path.suffix.lower() == ".docx":
        text = _docx_to_text(file_path)
        if text is not None:
            return text
    return file_path.read_bytes().decode("utf-8", errors="ignore")


def _docx_heading_level(style_name: str) -> int:
    """Map a Word paragraph style name to a markdown heading depth (0 = body)."""
    name = (style_name or "").strip()
    if name.lower() == "title":
        return 1
    m = _DOCX_HEADING_STYLE.match(name)
    if m:
        return min(6, max(1, int(m.group(1))))
    return 0


def _docx_to_text(file_path: Path) -> str | None:
    """Render a ``.docx`` to text in reading order, or ``None`` to fall back.

    Paragraphs and tables are walked via the document body's XML children so
    their original order is preserved (python-docx's ``paragraphs`` / ``tables``
    collections would otherwise separate them).
    """
    try:  # pragma: no cover - optional dependency
        import docx  # python-docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception:
        logger.warning("python-docx not installed; reading %s as raw text.", file_path.name)
        return None

    try:  # pragma: no cover - requires the dependency
        document = docx.Document(str(file_path))
        lines: list[str] = []
        for child in document.element.body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                rendered = _docx_paragraph_text(Paragraph(child, document))
                if rendered:
                    lines.append(rendered)
            elif tag == "tbl":
                rendered = _docx_table_text(Table(child, document))
                if rendered:
                    lines.append(rendered)
        return "\n".join(lines)
    except Exception as exc:  # pragma: no cover
        logger.warning("python-docx failed on %s (%s); reading as raw text.", file_path.name, exc)
        return None


def _docx_paragraph_text(para) -> str:  # pragma: no cover - requires the dependency
    text = para.text.strip()
    if not text:
        return ""
    style = para.style.name if para.style is not None else ""
    level = _docx_heading_level(style)
    return f"{'#' * level} {text}" if level else text


def _docx_table_text(table) -> str:  # pragma: no cover - requires the dependency
    rows = []
    for row in table.rows:
        cells = [" ".join((cell.text or "").split()) for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


# ---------------------------------------------------------------------------
# Header / footer suppression
# ---------------------------------------------------------------------------
def _normalise(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def _strip_boilerplate(blocks: list[_Block], num_pages: int) -> list[_Block]:
    """Drop repeated running headers/footers and bare page numbers.

    A short line that recurs near the top/bottom margin on a meaningful fraction
    of pages is boilerplate (e.g. a document code printed on every page). Removing
    it before sectioning is what lets a section that spans a page break stay whole
    instead of being split by the injected header line.
    """
    if num_pages <= 1:
        # Single page: only obvious page-number lines are safe to drop.
        return [b for b in blocks if not (not b.is_table and _PAGE_NUMBER.match(b.text.strip()))]

    pages_by_text: dict[str, set[int]] = defaultdict(set)
    for b in blocks:
        if not b.is_table and len(b.text) <= 90:
            pages_by_text[_normalise(b.text)].add(b.page)

    threshold = max(2, math.ceil(0.3 * num_pages))
    repeated = {norm for norm, pages in pages_by_text.items() if len(pages) >= threshold}

    kept: list[_Block] = []
    for b in blocks:
        text = b.text.strip()
        if not b.is_table and _PAGE_NUMBER.match(text):
            continue
        if not b.is_table and _normalise(b.text) in repeated and _in_margin(b):
            continue
        kept.append(b)
    return kept


def _in_margin(b: _Block) -> bool:
    """True when the line sits in the top or bottom 14% of the page."""
    if not b.page_height:
        return True
    rel = b.y0 / b.page_height
    return rel <= 0.14 or rel >= 0.86


# ---------------------------------------------------------------------------
# Heading classification
# ---------------------------------------------------------------------------
def _dominant_size(blocks: list[_Block]) -> float:
    """Most common font size across body lines, weighted by text length."""
    weights: Counter[float] = Counter()
    for b in blocks:
        if not b.is_table and b.size:
            weights[b.size] += len(b.text)
    return weights.most_common(1)[0][0] if weights else 0.0


def _classify_heading(text: str, size: float, bold: bool, body_size: float) -> tuple[int, str] | None:
    """Return ``(level, clean_heading)`` if the line is a heading, else None."""
    s = text.strip()
    if not s or _FIGURE_MARKER.match(s) or _PAGE_NUMBER.match(s):
        return None
    # Pipe-delimited rows are table content (rendered by the docx/PDF-text
    # fallback), never headings.
    if " | " in s:
        return None

    m = _MARKDOWN_HEADING.match(s)
    if m:
        return len(m.group(1)), m.group(2).strip()

    words = s.split()
    alpha = any(c.isalpha() for c in s)
    ends_sentence = s.endswith(_SENTENCE_END)

    m = _NUMBERED_HEADING.match(s)
    if m and alpha:
        number, rest = m.group(1), m.group(2).strip()
        depth = number.count(".") + 1
        if rest and not rest.endswith(_SENTENCE_END) and len(rest.split()) <= 10 and depth <= 3 and len(s) <= 80:
            return min(3, depth), rest

    # Typographic signal (PDF): clearly larger than body text.
    if body_size and size >= body_size + 1.0 and len(s) <= 120 and not ends_sentence:
        return (1 if size >= body_size + 3.0 else 2), s

    if not alpha or ends_sentence or len(words) > 12:
        return None

    # All-caps short line (reliable in standards/spec documents).
    if s == s.upper() and len(s) <= 70:
        return 2, s

    # Bold short line (PDF only).
    if bold and len(s) <= 80:
        return 2, s

    return None


def _classify(blocks: list[_Block], body_size: float) -> None:
    for b in blocks:
        if b.is_table:
            continue
        detected = _classify_heading(b.text, b.size, b.bold, body_size)
        if detected is not None:
            b.is_heading, b.level = True, detected[0]
            b.text = detected[1]


def _reassemble_headings(blocks: list[_Block]) -> list[_Block]:
    """Merge consecutive heading lines (a heading the PDF wrapped across lines).

    Only typographic (PDF) blocks carry a font ``size``; line-based sources
    (txt/md/docx) emit one complete heading per line, so merging adjacent
    headings there would wrongly fuse distinct sections — they are left as-is.
    """
    out: list[_Block] = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if not (b.is_heading and not b.is_table and b.size):
            out.append(b)
            i += 1
            continue
        text = b.text
        j = i + 1
        while (
            j < len(blocks)
            and blocks[j].is_heading
            and not blocks[j].is_table
            and blocks[j].level == b.level
            and blocks[j].page == b.page
            and len(text) + len(blocks[j].text) + 1 <= 160
        ):
            text = f"{text} {blocks[j].text}".strip()
            j += 1
        merged = _Block(
            text=text, page=b.page, page_height=b.page_height, y0=b.y0,
            size=b.size, bold=b.bold, is_heading=True, level=b.level,
        )
        out.append(merged)
        i = j
    return out


# ---------------------------------------------------------------------------
# Section assembly
# ---------------------------------------------------------------------------
def _assemble_full_text(blocks: list[_Block]) -> tuple[str, list[tuple[int, int, _Block]]]:
    parts: list[str] = []
    spans: list[tuple[int, int, _Block]] = []
    cursor = 0
    for b in blocks:
        segment = b.text if b.text.endswith("\n") else b.text + "\n"
        spans.append((cursor, cursor + len(segment), b))
        parts.append(segment)
        cursor += len(segment)
    return "".join(parts), spans


def _section_from_span(
    start: int, end: int, full_text: str, heading: str, level: int,
    page: int, spans: list[tuple[int, int, _Block]],
) -> DocumentSection | None:
    text = full_text[start:end]
    if not text.strip():
        return None
    contains_table = any(b.is_table for (s, _e, b) in spans if start <= s < end)
    return DocumentSection(
        id=str(uuid.uuid4()),
        page_number=page,
        heading=heading[:120] or "Section",
        heading_level=level,
        start_offset=start,
        end_offset=end,
        text=text,
        is_table=contains_table,
    )


def _sections_from_blocks(
    full_text: str, spans: list[tuple[int, int, _Block]]
) -> list[DocumentSection]:
    heading_positions = [k for k, (_s, _e, b) in enumerate(spans) if b.is_heading]
    if len(heading_positions) < 2:
        return _windowed_sections(full_text)

    sections: list[DocumentSection] = []
    first_start = spans[heading_positions[0]][0]
    if first_start > 0 and full_text[:first_start].strip():
        preamble = _section_from_span(
            0, first_start, full_text, "Introduction", 1, spans[0][2].page, spans
        )
        if preamble:
            sections.append(preamble)

    for pos, k in enumerate(heading_positions):
        start = spans[k][0]
        end = spans[heading_positions[pos + 1]][0] if pos + 1 < len(heading_positions) else len(full_text)
        b = spans[k][2]
        section = _section_from_span(start, end, full_text, b.text, b.level, b.page, spans)
        if section:
            sections.append(section)
    return sections


def _windowed_sections(full_text: str) -> list[DocumentSection]:
    """Fallback for unstructured documents with no detectable headings.

    Splits on blank-line paragraph boundaries, accumulating paragraphs into
    windows of roughly ``parent_max_tokens`` so downstream agents still receive
    coherent, offset-anchored sections (never used as the final chunking).
    """
    sections: list[DocumentSection] = []
    para_pattern = re.compile(r"\n\s*\n")
    cursor = 0
    win_start = 0
    win_tokens = 0
    target = SETTINGS.parent_max_tokens
    length = len(full_text)

    boundaries = [m.end() for m in para_pattern.finditer(full_text)] + [length]
    for boundary in boundaries:
        para = full_text[cursor:boundary]
        win_tokens += count_tokens(para)
        cursor = boundary
        if win_tokens >= target or boundary == length:
            text = full_text[win_start:boundary]
            if text.strip():
                heading = (text.strip().splitlines()[0] or "Section")[:80]
                sections.append(
                    DocumentSection(
                        id=str(uuid.uuid4()),
                        page_number=1,
                        heading=heading,
                        heading_level=1,
                        start_offset=win_start,
                        end_offset=boundary,
                        text=text,
                    )
                )
            win_start = boundary
            win_tokens = 0
    return sections


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def parse_document(file_path: Path) -> tuple[str, list[DocumentSection]]:
    """Parse a document into ``(full_text, sections)`` with preserved offsets."""
    extracted = None
    if file_path.suffix.lower() == ".pdf":
        extracted = _extract_pdf_blocks(file_path)

    if extracted is not None:
        blocks, num_pages = extracted
    else:
        raw = _extract_text(file_path)
        if not raw.strip():
            return raw, []
        blocks, num_pages = _blocks_from_text(raw), 1

    if not blocks:
        return "", []

    blocks = _strip_boilerplate(blocks, num_pages)
    if not blocks:
        return "", []

    body_size = _dominant_size(blocks)
    _classify(blocks, body_size)
    blocks = _reassemble_headings(blocks)

    full_text, spans = _assemble_full_text(blocks)
    sections = _sections_from_blocks(full_text, spans)

    num_headings = sum(1 for s in sections if s.heading not in ("Introduction", "Section"))
    logger.info(
        "Parsed %s: %d page(s) -> %d section(s) (%d heading-anchored, %d table).",
        file_path.name, num_pages, len(sections), num_headings,
        sum(1 for s in sections if s.is_table),
    )
    return full_text, sections
